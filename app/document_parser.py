"""Парсер документів: PDF, EPUB, FB2, MOBI/AZW3, DOCX, ODT, RTF, HTML → текст."""

import io
import os
import logging
import re
from typing import Iterable, Optional

logger = logging.getLogger(__name__)


def _html_to_plain(html: str) -> str:
    """HTML → звичайний текст через lxml (зберігає абзаци)."""
    from lxml import html as lhtml, etree

    # прибираємо XML-декларацію (не підтримується lxml.html.fromstring з рядком)
    cleaned = re.sub(r"^\s*<\?xml[^>]*\?>", "", html, flags=re.IGNORECASE)
    tree = lhtml.fromstring(cleaned)
    # вставляємо роздільники між блочними елементами
    for br in tree.xpath("//br"):
        br.tail = (br.tail or "") + "\n"
    for p in tree.xpath("//p | //div | //h1 | //h2 | //h3 | //li"):
        text = "".join(p.itertext())
        if text:
            p.text = text
            p.tail = (p.tail or "") + "\n\n"
    text = etree.tostring(tree, method="text", encoding="unicode")
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def parse_pdf(data: bytes) -> str:
    """Текст з PDF (pypdf, підтримка зашифрованих, битих сторінок і текстового сміття)."""
    try:
        from pypdf import PdfReader
    except ImportError:
        from PyPDF2 import PdfReader  # type: ignore

    reader = PdfReader(io.BytesIO(data))
    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            pass

    parts = []
    for page in reader.pages:
        try:
            text = page.extract_text()
        except Exception:
            text = ""
        if text and text.strip():
            parts.append(text.strip())
    return "\n\n".join(parts)


def parse_epub(data: bytes) -> str:
    """Текст з EPUB (порядок глав, заголовки голов, пропуск битих)."""
    from ebooklib import epub
    from lxml import etree

    book = epub.read_epub(io.BytesIO(data), options={"ignore_ncx": True, "ignore_smil": True})
    doc_type = getattr(epub, "ITEM_DOCUMENT", 9)
    parts = []

    for item in book.get_items():
        if item.get_type() != doc_type:
            continue
        title = item.get_name()
        if title and title.lower().lstrip("/").startswith("nav"):
            continue
        html = item.get_content()
        try:
            text = _html_to_plain(html.decode("utf-8", errors="replace"))
        except Exception:
            text = ""
        if not text:
            continue
        parts.append(text)

    return "\n\n".join(parts)


def parse_fb2(data: bytes) -> str:
    """Текст з FB2 (FictionBook2, XML)."""
    from lxml import etree

    tree = etree.parse(io.BytesIO(data))
    ns = {"fb": "http://www.gribuser.ru/xml/fictionbook/2.0"}

    parts = []
    for body in tree.xpath("//fb:body", namespaces=ns):
        for elem in body.iter():
            tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ""
            if tag in ("p", "subtitle", "title", "epigraph", "cite", "text-author"):
                text = etree.tostring(elem, method="text", encoding="unicode").strip()
                if text:
                    parts.append(text)
            elif tag == "empty-line":
                parts.append("")
    return "\n".join(parts)


def parse_mobi(data: bytes) -> str:
    """Текст з MOBI/AZW3 (через тимчасову папку)."""
    import tempfile
    import shutil
    import mobi

    tmp_dir = tempfile.mkdtemp()
    try:
        tmp_path = os.path.join(tmp_dir, "book.mobi")
        with open(tmp_path, "wb") as f:
            f.write(data)

        result = mobi.extract(tmp_path)
        opf_path = result.get("opf")
        if not opf_path:
            html_path = result.get("html")
            if not html_path:
                return ""
            with open(html_path, "r", encoding="utf-8", errors="replace") as f:
                html = f.read()
            return _html_to_plain(html)

        from lxml import etree
        opf_dir = os.path.dirname(opf_path)
        with open(opf_path, "r", encoding="utf-8", errors="replace") as f:
            opf = f.read()
        tree = etree.parse(io.BytesIO(opf.encode()))
        ns = {"opf": "http://www.idpf.org/2007/opf"}
        spine = tree.xpath("//opf:spine/opf:itemref/@idref", namespaces=ns)
        manifest = tree.xpath("//opf:manifest/opf:item", namespaces=ns)
        id_to_href = {
            item.get("id"): item.get("href")
            for item in manifest
            if item.get("id") and item.get("href")
        }

        parts = []
        for item_id in spine:
            href = id_to_href.get(item_id)
            if not href:
                continue
            fpath = os.path.join(opf_dir, href)
            if os.path.exists(fpath):
                with open(fpath, "r", encoding="utf-8", errors="replace") as f:
                    html = f.read()
                text = _html_to_plain(html)
                if text:
                    parts.append(text)
        return "\n\n".join(parts)
    except Exception as e:
        logger.error("MOBI parse error: %s", e)
        raise
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def parse_docx(data: bytes) -> str:
    """Текст з DOCX (python-docx, абзаци й таблиці)."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n\n".join(parts)


def parse_odt(data: bytes) -> str:
    """Текст з ODT (odfpy)."""
    from odf.opendocument import load
    from odf import teletype

    doc = load(io.BytesIO(data))
    parts = []

    def walk(elem):
        for child in elem.childNodes:
            qname = getattr(child, "qname", None)
            if isinstance(qname, tuple) and qname[1] in ("p", "h"):
                text = teletype.extractText(child).strip()
                if text:
                    parts.append(text)
            walk(child)

    walk(doc.body)
    return "\n\n".join(parts)


_ANSI = "cp1252"


def parse_rtf(data: bytes) -> str:
    """Текст з RTF (strip-парс — без зовнішніх залежностей)."""
    try:
        text = data.decode("utf-8", errors="strict")
    except UnicodeDecodeError:
        text = data.decode(_ANSI, errors="replace")

    return _strip_rtf(text)


def _strip_rtf(text: str) -> str:
    """Видаляє RTF-керуючі коди, залишає чистий текст."""
    # підсунути екранування
    text = text.replace("\\par", "\n").replace("\\line", "\n")
    text = text.replace("\\tab", "\t")
    # видалити групи керуючих слів і коди
    text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)
    # спецсимволи
    text = text.replace("\\{", "{").replace("\\}", "}").replace("\\\\", "\\")
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def parse_html(data: bytes) -> str:
    """Текст з HTML."""
    return _html_to_plain(data.decode("utf-8", errors="replace"))


PARSERS = {
    ".pdf": parse_pdf,
    ".epub": parse_epub,
    ".fb2": parse_fb2,
    ".mobi": parse_mobi,
    ".azw3": parse_mobi,
    ".docx": parse_docx,
    ".odt": parse_odt,
    ".rtf": parse_rtf,
    ".html": parse_html,
    ".htm": parse_html,
}


def parse_document(filename: str, data: bytes) -> tuple[str, str]:
    """Розпізнає формат за розширенням та повертає (текст, формат).

    Повертає (порожній рядок, "") якщо формат не підтримується.
    """
    ext = os.path.splitext(filename)[1].lower()
    parser = PARSERS.get(ext)
    if not parser:
        return "", ""
    try:
        text = parser(data)
        fmt = ext.lstrip(".")
        return text, fmt
    except Exception as e:
        logger.error("Parse error for %s: %s", filename, e)
        raise
